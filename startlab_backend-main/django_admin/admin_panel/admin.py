from django.contrib import admin
from django_summernote.admin import SummernoteModelAdmin
from django.urls import path, reverse
from django.shortcuts import redirect
from django.db.models import Count, Min, Max
from django.http import HttpResponse
from django.utils.html import format_html
from django.utils.safestring import mark_safe
from django.core.paginator import Paginator
from django.db import connection
from django.utils.encoding import smart_str
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from urllib.parse import quote
import re
from django.utils.text import slugify

from .models import Partner, FAQ, News, Answer, Answer2, Team, Supervisor2


# =========================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# =========================

# Границы ячеек (совместимость с iPhone/Safari)
def set_cell_border(cell, **kwargs):
    """
    Устанавливает рамки ячейки.
    Пример:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000"},
        bottom={"sz": 12, "val": "single", "color": "000000"},
        left={"sz": 12, "val": "single", "color": "000000"},
        right={"sz": 12, "val": "single", "color": "000000"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            element = tcBorders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tcBorders.append(element)
            for key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_table_borders(table):
    """Устанавливает чёрные границы для всех ячеек таблицы"""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            )


def _safe_filename(title: str, ext: str = "docx"):
    """
    Возвращает кортеж (ascii_fallback, utf8_filename) без запрещённых символов.
    Safari/iOS ориентируется на filename*, другие — на filename.
    """
    title = (title or "document").strip()
    # Уберём запрещённые символы для имён файлов и лишние пробелы
    cleaned = re.sub(r'[\\/*?:"<>|\r\n\t]', "_", title)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    # Ограничим длину
    cleaned = cleaned[:150]

    ascii_fallback_base = slugify(cleaned) or "document"
    ascii_fallback = f"{ascii_fallback_base}.{ext}"
    utf8_filename = f"{cleaned}.{ext}"
    return ascii_fallback, utf8_filename


def _stream_docx_response(doc: Document, title: str):
    """
    Собирает HttpResponse для DOCX с корректными заголовками Content-Disposition:
    - filename (ASCII фолбэк)
    - filename* (RFC 5987, UTF-8) — важно для iOS/Safari
    """
    ascii_fallback, utf8_filename = _safe_filename(title, "docx")
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = (
        f"attachment; filename={ascii_fallback}; filename*=UTF-8''{quote(utf8_filename)}"
    )
    response["X-Content-Type-Options"] = "nosniff"
    doc.save(response)
    return response


# =========================
# PARTNER
# =========================

@admin.register(Partner)
class PartnerAdmin(SummernoteModelAdmin):
    """Админка для партнеров"""
    list_display = ('name', 'title', 'website', 'is_active', 'created_at', 'logo_preview')
    list_filter = ('is_active', 'created_at')
    search_fields = ('name', 'description')
    ordering = ('name',)
    list_editable = ('is_active',)

    fieldsets = (
        ('Основная информация', {
            'fields': ('name', 'title', 'logo', 'logo_preview', 'description', 'website', 'is_active')
        }),
        ('Метаданные', {
            'fields': ('created_at',),
            'classes': ('collapse',)
        }),
    )
    readonly_fields = ('created_at', 'logo_preview')
    summernote_fields = ('title', 'description')

    def logo_preview(self, obj):
        """Показывает превью логотипа"""
        if obj.logo:
            return format_html('<img src="{}" style="max-height:100px;max-width:100px;" />', obj.logo.url)
        return "Нет логотипа"
    logo_preview.short_description = "Превью логотипа"


# =========================
# FAQ
# =========================

@admin.register(FAQ)
class FAQAdmin(SummernoteModelAdmin):
    """Админка для FAQ"""
    list_display = ('question', 'order', 'is_active', 'created_at', 'image_preview')
    list_filter = ('is_active', 'created_at')
    search_fields = ('question', 'answer')
    ordering = ('order',)
    list_editable = ('order', 'is_active')

    fieldsets = (
        ('Основная информация', {
            'fields': ('question', 'answer', 'image', 'image_preview', 'order', 'is_active')
        }),
        ('Метаданные', {
            'fields': ('created_at',),
            'classes': ('collapse',)
        }),
    )
    readonly_fields = ('created_at', 'image_preview')
    summernote_fields = ('question', 'answer')

    def image_preview(self, obj):
        """Показывает превью изображения"""
        if obj.image:
            return format_html('<img src="{}" style="max-height:100px;max-width:100px;" />', obj.image.url)
        return "Нет изображения"
    image_preview.short_description = "Превью изображения"


# =========================
# NEWS
# =========================

@admin.register(News)
class NewsAdmin(SummernoteModelAdmin):
    """Админка для новостей"""
    list_display = ('title', 'is_active', 'created_at', 'updated_at', 'image_preview')
    list_filter = ('is_active', 'created_at')
    search_fields = ('title', 'content')
    ordering = ('-created_at',)
    list_editable = ('is_active',)

    fieldsets = (
        ('Основная информация', {
            'fields': ('title', 'content', 'image', 'image_preview', 'is_active')
        }),
        ('Метаданные', {
            'fields': ('created_at', 'updated_at'),
            'classes': ('collapse',)
        }),
    )
    readonly_fields = ('created_at', 'updated_at', 'image_preview')
    summernote_fields = ('content',)

    def image_preview(self, obj):
        """Показывает превью изображения"""
        if obj.image:
            return format_html('<img src="{}" style="max-height:100px;max-width:100px;" />', obj.image.url)
        return "Нет изображения"
    image_preview.short_description = "Превью изображения"


# =========================
# SCIENCE (Answer)
# =========================

def _add_row_science(table, title, value):
    row = table.add_row()
    row.cells[0].text = title
    row.cells[1].text = smart_str(value or "—")


def _add_heading_science(doc, text):
    try:
        doc.add_paragraph(text, style="Heading 2")
    except Exception:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True


@admin.register(Answer)
class AnswerAdmin(admin.ModelAdmin):
    list_display = ("title", "budgetBYN", "timeline", "sha256")
    search_fields = ("title", "relevance", "description", "sha256")
    list_filter = ("budgetBYN",)
    readonly_fields = ("sha256", "download_docx_button", "participants_and_supervisor")
    actions = ["export_docx"]
    fields = ("download_docx_button", "title", "relevance", "goal", "tasks", "description", "expectedResults",
              "marketAssessment", "competitionAnalysis", "budgetBYN", "timeline", "label", "url",
              "additionalInfo", "sha256", "participants_and_supervisor")

    def get_urls(self):
        urls = super().get_urls()
        custom_urls = [
            path('<int:object_id>/export-docx/', self.admin_site.admin_view(self.export_docx_view),
                 name='admin_answer_export_docx'),
        ]
        return custom_urls + urls

    def download_docx_button(self, obj):
        if not obj or not obj.id:
            return "—"
        url = reverse('admin:admin_answer_export_docx', args=[obj.id])
        return format_html('<a class="button" href="{}">Скачать DOCX (по шаблону)</a>', url)
    download_docx_button.short_description = "Экспорт DOCX"

    def export_docx(self, request, queryset):
        if queryset.count() != 1:
            self.message_user(request, "Выберите ровно одну запись для экспорта", level=admin.messages.WARNING)
            return
        obj = queryset.first()
        return self._generate_docx(obj)
    export_docx.short_description = "Сформировать DOCX (наука)"

    def export_docx_view(self, request, object_id):
        obj = self.get_object(request, object_id)
        if not obj:
            return HttpResponse(status=404)
        return self._generate_docx(obj)

    def _generate_docx(self, obj):
        doc = Document()

        # Базовый стиль
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.left_indent = Pt(20)
        style.paragraph_format.space_after = Pt(6)

        # Шапка
        def add_centered(text, size=14):
            p = doc.add_paragraph(text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.runs[0]
            run.font.size = Pt(size)
            run.font.name = "Times New Roman"

        add_centered("ЗАЯВКА", size=14)
        add_centered("на участие в Конкурсном отборе научных и инновационных идей обучающихся", size=14)
        doc.add_paragraph("")
        add_centered("(СТАРТАПОВ и НАУЧНЫХ ПРОЕКТОВ)", size=14)
        doc.add_paragraph("")

        # Основная таблица
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        _add_row_science(table, "Наименование научной или инновационной идеи проекта / стартапа", obj.title)
        _add_row_science(table, "Актуальность научной или инновационной идеи", obj.relevance)
        _add_row_science(table, "Цель научной или инновационной идеи", obj.goal)
        _add_row_science(table, "Задачи научной или инновационной идеи", obj.tasks)
        _add_row_science(table, "Краткое описание научной или инновационной идеи (не более 3000 символов)", obj.description)
        _add_row_science(table, "Ожидаемые результаты научной или инновационной идеи", obj.expectedResults)
        _add_row_science(table, "Оценка потенциального рынка", obj.marketAssessment)
        _add_row_science(table, "Конкурентный анализ", obj.competitionAnalysis)
        _add_row_science(table, "Предполагаемый бюджет (бел. руб.)", obj.budgetBYN)
        _add_row_science(table, "Предполагаемый срок реализации научной или инновационной идеи", obj.timeline)

        set_table_borders(table)

        # Участники команды
        members = Team.objects.filter(sha256=obj.sha256)
        if members.exists():
            participant_fields = [
                "ФИО",
                "Факультет",
                "Группа",
                "Мобильный телефон",
                "Электронная почта",
                "Ключевые компетенции участника",
                "Роль в реализации научной или инновационной идее"
            ]

            table_members = doc.add_table(rows=len(participant_fields) + 1, cols=members.count() + 1)
            table_members.style = "Table Grid"

            table_members.cell(0, 0).text = "Участники команды"
            for col_idx, m in enumerate(members, start=1):
                table_members.cell(0, col_idx).text = f"Участник {col_idx}"

            for row_idx, field_name in enumerate(participant_fields, start=1):
                table_members.cell(row_idx, 0).text = field_name
                for col_idx, m in enumerate(members, start=1):
                    value = "—"
                    if field_name == "ФИО":
                        name_parts = []
                        if m.lastName:
                            name_parts.append(m.lastName.strip())
                        if m.firstName:
                            name_parts.append(m.firstName.strip())
                        if m.middleName:
                            name_parts.append(m.middleName.strip())
                        value = " ".join(name_parts) if name_parts else "—"
                    elif field_name == "Факультет":
                        value = m.faculty or "—"
                    elif field_name == "Группа":
                        value = m.group or "—"
                    elif field_name == "Мобильный телефон":
                        value = m.phone or "—"
                    elif field_name == "Электронная почта":
                        value = m.email or "—"
                    elif field_name == "Ключевые компетенции участника":
                        value = m.keySkills or "—"
                    elif field_name == "Роль в реализации научной или инновационной идее":
                        value = m.role or "—"
                    table_members.cell(row_idx, col_idx).text = value

            set_table_borders(table_members)
        else:
            doc.add_paragraph("Участники команды: —")

        # Научный руководитель
        sup = Supervisor2.objects.filter(sha256=obj.sha256).first()
        if sup:
            table_supervisor = doc.add_table(rows=5, cols=2)
            table_supervisor.style = "Table Grid"
            table_supervisor.cell(0, 0).text = "Научный руководитель"
            table_supervisor.cell(0, 1).text = ""
            table_supervisor.cell(1, 0).text = "ФИО"
            table_supervisor.cell(1, 1).text = sup.fullName or "—"
            table_supervisor.cell(2, 0).text = "Звание"
            table_supervisor.cell(2, 1).text = sup.academicTitle or "—"
            table_supervisor.cell(3, 0).text = "Должность"
            table_supervisor.cell(3, 1).text = sup.position or "—"
            table_supervisor.cell(4, 0).text = "Мобильный телефон"
            table_supervisor.cell(4, 1).text = sup.phone or "—"
            set_table_borders(table_supervisor)
        else:
            doc.add_paragraph("Научный руководитель: —")

        # Приложения/доп.информация
        table_info = doc.add_table(rows=4, cols=1)
        table_info.style = "Table Grid"
        table_info.cell(0, 0).text = "Приложение"

        table_info.cell(1, 0).text = obj.url or "—"
        table_info.cell(2, 0).text = "Дополнительная информация"
        table_info.cell(3, 0).text = obj.additionalInfo or "—"
        set_table_borders(table_info)

        # Отдаём с корректными заголовками имени файла
        return _stream_docx_response(doc, obj.title)

    def participants_and_supervisor(self, obj):
        """HTML: участники и (если есть) научный руководитель по sha256"""
        if not obj:
            return "—"
        parts = []
        members = list(Team.objects.filter(sha256=obj.sha256))
        if members:
            parts.append('<h3>Участники команды</h3>')
            parts.append('<table class="table" style="width:100%;border-collapse:collapse;">')
            parts.append('<thead>\n<tr>\n'
                         '<th style="border:1px solid #ccc;padding:4px;">ФИО</th>'
                         '<th style="border:1px solid #ccc;padding:4px;">Факультет</th>'
                         '<th style="border:1px solid #ccc;padding:4px;">Группа</th>'
                         '<th style="border:1px solid #ccc;padding:4px;">Телефон</th>'
                         '<th style="border:1px solid #ccc;padding:4px;">Email</th>'
                         '<th style="border:1px solid #ccc;padding:4px;">Ключевые навыки</th>'
                         '<th style="border:1px solid #ccc;padding:4px;">Роль</th>'
                         '</tr>\n</thead><tbody>')
            for m in members:
                full_name = " ".join(filter(None, [
                    (m.lastName or "").strip(),
                    (m.firstName or "").strip(),
                    (m.middleName or "").strip(),
                ])).strip() or "—"
                team_url = reverse('admin:admin_panel_team_change', args=[m.id]) if m.id else None
                link_html = f'<a href="{team_url}">{full_name}</a>' if team_url else full_name
                parts.append(
                    '<tr>'
                    f'<td style="border:1px solid #eee;padding:4px;">{link_html}</td>'
                    f'<td style="border:1px solid #eee;padding:4px;">{m.faculty or "—"}</td>'
                    f'<td style="border:1px solid #eee;padding:4px;">{m.group or "—"}</td>'
                    f'<td style="border:1px solid #eee;padding:4px;">{m.phone or "—"}</td>'
                    f'<td style="border:1px solid #eee;padding:4px;">{m.email or "—"}</td>'
                    f'<td style="border:1px solid #eee;padding:4px;">{m.keySkills or "—"}</td>'
                    f'<td style="border:1px solid #eee;padding:4px;">{m.role or "—"}</td>'
                    '</tr>'
                )
            parts.append('</tbody></table>')
        else:
            parts.append('<p><em>Участники не найдены</em></p>')

        supervisor = Supervisor2.objects.filter(sha256=obj.sha256).first()
        parts.append('<h3>Научный руководитель</h3>')
        if supervisor:
            parts.append('<table class="table" style="width:100%;border-collapse:collapse;">')
            parts.append('<tbody>')
            parts.append(
                '<tr>'
                '<th style="text-align:left;border:1px solid #ccc;padding:4px;width:220px;">ФИО</th>'
                f'<td style="border:1px solid #eee;padding:4px;">{supervisor.fullName or "—"}</td>'
                '</tr>'
            )
            parts.append(
                '<tr>'
                '<th style="text-align:left;border:1px solid #ccc;padding:4px;">Звание</th>'
                f'<td style="border:1px solid #eee;padding:4px;">{supervisor.academicTitle or "—"}</td>'
                '</tr>'
            )
            parts.append(
                '<tr>'
                '<th style="text-align:left;border:1px solid #ccc;padding:4px;">Должность</th>'
                f'<td style="border:1px solid #eee;padding:4px;">{supervisor.position or "—"}</td>'
                '</tr>'
            )
            parts.append(
                '<tr>'
                '<th style="text-align:left;border:1px solid #ccc;padding:4px;">Телефон</th>'
                f'<td style="border:1px solid #eee;padding:4px;">{supervisor.phone or "—"}</td>'
                '</tr>'
            )
            parts.append(
                '<tr>'
                '<th style="text-align:left;border:1px solid #ccc;padding:4px;">Email</th>'
                f'<td style="border:1px solid #eee;padding:4px;">{supervisor.email or "—"}</td>'
                '</tr>'
            )
            parts.append('</tbody></table>')
        else:
            parts.append('<p><em>Руководитель не указан</em></p>')
        return mark_safe("".join(parts))
    participants_and_supervisor.short_description = "Участники и руководитель"


# =========================
# STARTUP (Answer2)
# =========================

def _add_row_startup(table, key, value):
    cells = table.add_row().cells
    cells[0].text = key or "—"
    cells[1].text = str(value or "—")
    # фиксируем ширину столбцов
    cells[0].width = Cm(8)
    cells[1].width = Cm(10)


def _add_heading_startup(doc, text):
    doc.add_paragraph()
    doc.add_heading(text, level=2)


@admin.register(Answer2)
class Answer2Admin(admin.ModelAdmin):
    list_display = ("title", "budgetBYN", "timeline", "sha256")
    search_fields = ("title", "description", "sha256")
    list_filter = ("budgetBYN",)
    readonly_fields = ("sha256", "download_docx_button", "participants_overview")
    actions = ["export_docx"]
    fields = (
        "download_docx_button", "title", "problemStatementShort", "goal", "stageAndNextSteps",
        "description", "founderMotivationAndExpertise", "expectedResults", "benefitForBelarus",
        "marketAssessment", "monetization", "competitionAnalysis", "budgetBYN", "needsInvestmentNow",
        "timeline", "label", "url", "additionalInfo", "sha256", "participants_overview"
    )

    def get_urls(self):
        urls = super().get_urls()
        custom_urls = [
            path('<int:object_id>/export-docx/', self.admin_site.admin_view(self.export_docx_view),
                 name='admin_answer2_export_docx'),
        ]
        return custom_urls + urls

    def download_docx_button(self, obj):
        if not obj or not obj.id:
            return "—"
        url = reverse('admin:admin_answer2_export_docx', args=[obj.id])
        return format_html('<a class="button" href="{}">Скачать DOCX (по шаблону)</a>', url)
    download_docx_button.short_description = "Экспорт DOCX"

    def export_docx(self, request, queryset):
        if queryset.count() != 1:
            self.message_user(request, "Выберите ровно одну запись для экспорта", level=admin.messages.WARNING)
            return
        obj = queryset.first()
        return self._generate_docx(obj)
    export_docx.short_description = "Сформировать DOCX (стартап)"

    def export_docx_view(self, request, object_id):
        obj = self.get_object(request, object_id)
        if not obj:
            return HttpResponse(status=404)
        return self._generate_docx(obj)

    def _generate_docx(self, obj):
        doc = Document()

        # Базовый стиль
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.left_indent = Pt(20)
        style.paragraph_format.space_after = Pt(6)

        # Шапка
        def add_centered(text, size=14):
            p = doc.add_paragraph(text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.runs[0]
            run.font.size = Pt(size)
            run.font.name = "Times New Roman"

        add_centered("ЗАЯВКА", size=14)
        add_centered("на участие в Конкурсном отборе научных и инновационных идей обучающихся", size=14)
        doc.add_paragraph("")
        add_centered("(СТАРТАПОВ и НАУЧНЫХ ПРОЕКТОВ)", size=14)
        doc.add_paragraph("")

        # Основная таблица
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        _add_row_startup(table, "Наименование научной или инновационной идеи проекта / стартапа", obj.title)
        _add_row_startup(table, "Актуальность научной или инновационной идеи", obj.problemStatementShort)
        _add_row_startup(table, "Цель научной или инновационной идеи", obj.goal)
        _add_row_startup(table, "Задачи научной или инновационной идеи", obj.stageAndNextSteps)
        _add_row_startup(table, "Краткое описание научной или инновационной идеи (не более 3000 символов)", obj.problemStatementShort)
        _add_row_startup(table, "Ожидаемые результаты научной или инновационной идеи", obj.benefitForBelarus)
        _add_row_startup(table, "Оценка потенциального рынка", obj.monetization)
        _add_row_startup(table, "Конкурентный анализ", obj.competitionAnalysis)
        _add_row_startup(table, "Предполагаемый бюджет (бел. руб.)", obj.needsInvestmentNow)
        _add_row_startup(table, "Предполагаемый срок реализации научной или инновационной идеи", obj.timeline)

        set_table_borders(table)

        # Участники команды
        members = Team.objects.filter(sha256=obj.sha256)
        if members.exists():
            participant_fields = [
                "ФИО",
                "Факультет",
                "Группа",
                "Мобильный телефон",
                "Электронная почта",
                "Ключевые компетенции участника",
                "Роль в реализации научной или инновационной идее"
            ]

            table_members = doc.add_table(rows=len(participant_fields) + 1, cols=members.count() + 1)
            table_members.style = "Table Grid"

            table_members.cell(0, 0).text = "Участники команды"
            for col_idx, m in enumerate(members, start=1):
                table_members.cell(0, col_idx).text = f"Участник {col_idx}"

            for row_idx, field_name in enumerate(participant_fields, start=1):
                table_members.cell(row_idx, 0).text = field_name
                for col_idx, m in enumerate(members, start=1):
                    value = "—"
                    if field_name == "ФИО":
                        name_parts = []
                        if m.lastName:
                            name_parts.append(m.lastName.strip())
                        if m.firstName:
                            name_parts.append(m.firstName.strip())
                        if m.middleName:
                            name_parts.append(m.middleName.strip())
                        value = " ".join(name_parts) if name_parts else "—"
                    elif field_name == "Факультет":
                        value = m.faculty or "—"
                    elif field_name == "Группа":
                        value = m.group or "—"
                    elif field_name == "Мобильный телефон":
                        value = m.phone or "—"
                    elif field_name == "Электронная почта":
                        value = m.email or "—"
                    elif field_name == "Ключевые компетенции участника":
                        value = m.keySkills or "—"
                    elif field_name == "Роль в реализации научной или инновационной идее":
                        value = m.role or "—"
                    table_members.cell(row_idx, col_idx).text = value

            set_table_borders(table_members)
        else:
            doc.add_paragraph("Участники команды: —")

        # Приложения/доп.информация
        table_info = doc.add_table(rows=4, cols=1)
        table_info.style = "Table Grid"
        table_info.cell(0, 0).text = "Приложение"
        table_info.cell(1, 0).text = obj.url or "—"
        table_info.cell(2, 0).text = "Дополнительная информация"
        table_info.cell(3, 0).text = obj.additionalInfo or "—"
        set_table_borders(table_info)

        # Отдаём с корректными заголовками имени файла
        return _stream_docx_response(doc, obj.title)

    def participants_overview(self, obj):
        """HTML: участники команды по sha256 (для стартапа)"""
        if not obj:
            return "—"
        members = list(Team.objects.filter(sha256=obj.sha256))
        if not members:
            return mark_safe('<p><em>Участники не найдены</em></p>')
        parts = []
        parts.append('<h3>Участники команды</h3>')
        parts.append('<table class="table" style="width:100%;border-collapse:collapse;">')
        parts.append('<thead>\n<tr>\n'
                     '<th style="border:1px solid #ccc;padding:4px;">ФИО</th>'
                     '<th style="border:1px solid #ccc;padding:4px;">Факультет</th>'
                     '<th style="border:1px solid #ccc;padding:4px;">Группа</th>'
                     '<th style="border:1px solid #ccc;padding:4px;">Телефон</th>'
                     '<th style="border:1px solid #ccc;padding:4px;">Email</th>'
                     '<th style="border:1px solid #ccc;padding:4px;">Ключевые навыки</th>'
                     '<th style="border:1px solid #ccc;padding:4px;">Роль</th>'
                     '</tr>\n</thead><tbody>')
        for m in members:
            full_name = " ".join(filter(None, [
                (m.lastName or "").strip(),
                (m.firstName or "").strip(),
                (m.middleName or "").strip(),
            ])).strip() or "—"
            team_url = reverse('admin:admin_panel_team_change', args=[m.id]) if m.id else None
            link_html = f'<a href="{team_url}">{full_name}</a>' if team_url else full_name
            parts.append(
                '<tr>'
                f'<td style="border:1px solid #eee;padding:4px;">{link_html}</td>'
                f'<td style="border:1px solid #eee;padding:4px;">{m.faculty or "—"}</td>'
                f'<td style="border:1px solid #eee;padding:4px;">{m.group or "—"}</td>'
                f'<td style="border:1px solid #eee;padding:4px;">{m.phone or "—"}</td>'
                f'<td style="border:1px solid #eee;padding:4px;">{m.email or "—"}</td>'
                f'<td style="border:1px solid #eee;padding:4px;">{m.keySkills or "—"}</td>'
                f'<td style="border:1px solid #eee;padding:4px;">{m.role or "—"}</td>'
                '</tr>'
            )
        parts.append('</tbody></table>')
        return mark_safe("".join(parts))
    participants_overview.short_description = "Участники команды"


# =========================
# TEAM
# =========================

@admin.register(Team)
class TeamAdmin(admin.ModelAdmin):
    list_display = ("lastName", "firstName", "middleName", "faculty", "group", "phone", "email", "role", "sha256")
    search_fields = ("lastName", "firstName", "middleName", "email", "phone", "sha256")
    list_filter = ("faculty", "group", "role")
    ordering = ("lastName", "firstName")

    fieldsets = (
        ('Личная информация', {
            'fields': ('lastName', 'firstName', 'middleName')
        }),
        ('Учебная информация', {
            'fields': ('faculty', 'group')
        }),
        ('Контактная информация', {
            'fields': ('phone', 'email')
        }),
        ('Профессиональная информация', {
            'fields': ('keySkills', 'role')
        }),
        ('Системная информация', {
            'fields': ('sha256',),
            'classes': ('collapse',)
        }),
    )

    readonly_fields = ('sha256',)

    def get_queryset(self, request):
        """Оптимизируем запросы"""
        return super().get_queryset(request)

    def save_model(self, request, obj, form, change):
        """Автоматически генерируем SHA256 при создании новой записи"""
        if not change:  # Если это новая запись
            import hashlib
            import json
            data = {
                'lastName': obj.lastName or '',
                'firstName': obj.firstName or '',
                'middleName': obj.middleName or '',
                'faculty': obj.faculty or '',
                'group': obj.group or '',
                'phone': obj.phone or '',
                'email': obj.email or '',
                'keySkills': obj.keySkills or '',
                'role': obj.role or ''
            }
            canonical = json.dumps(data, ensure_ascii=False, sort_keys=True)
            obj.sha256 = hashlib.sha256(canonical.encode("utf-8")).hexdigest()
        super().save_model(request, obj, form, change)


# =========================
# SUPERVISOR2
# =========================

@admin.register(Supervisor2)
class Supervisor2Admin(admin.ModelAdmin):
    list_display = ("fullName", "academicTitle", "position", "email", "sha256")
    search_fields = ("fullName", "email", "sha256")


